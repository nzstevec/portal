import pkg_resources
import pandas as pd
from pathlib import Path
from io import BytesIO, StringIO
from tempfile import NamedTemporaryFile
# import streamlit as st

import fitz
import docx
import docx2txt
import tiktoken
from striprtf.striprtf import rtf_to_text

from gpt.prompts import (
    summarization_prompt_rimon_specific,
    summarization_prompt_rimon_specific_new,
    summarization_prompt_rimon_specific_json_style,
)
from service.runpod_utils import runpod_call
from config import logger

NEW_TEMPLATE_PATH = pkg_resources.resource_filename(
    __name__, "Templates/BLANK - RDTI Initial Information Preparation FY23 - v18.docx"
)

OLD_TEMPLATE_PATH = pkg_resources.resource_filename(
    __name__, "Templates/BLANK - RDTI Initial Information Preparation FY22 - v14.docx"
)

TEMPLATE_PATH = pkg_resources.resource_filename(
    __name__, "Templates/rimon_template_v2.docx"
)


def count_tokens(input_text: str) -> int:
    """
    Count the number of tokens in an input string
    """
    encoding_model = tiktoken.get_encoding("cl100k_base")
    return len(encoding_model.encode(input_text))


class DocumentParser:
    def __init__(self, files: list = None, encoding_model: str = "cl100k_base"):
        self.type = None
        self.files = files
        self.encoding_model = tiktoken.get_encoding(encoding_model)

    def load_from_file_path(self, file_path) -> str:

        file_extension = Path(file_path).suffix

        match file_extension:
            case ".pdf":
                contents = self.load_pdf(file_path=file_path)
            case ".txt":
                contents = self.load_txt(file_path=file_path)
            case ".docx":
                contents = self.load_docx_advanced(file_path=file_path)
            case ".rtf":
                contents = self.load_rtf(file_path=file_path)
            case ".csv":
                contents = self.load_csv(file_path=file_path)
            case ".xlsx":
                contents = self.load_xlsx(file_path=file_path)
            case _:
                contents = "Unknown file type"

        return contents

    def count_tokens(self, input_text: str) -> int:
        """
        Count the number of tokens in an input string
        """
        return len(self.encoding_model.encode(input_text))

    def load_file_contents(self) -> str:
        """
        Takes the files passed to the document parser, and returns a string with all of their contents.
        """
        TEMPLATE_FLAG = False
        rimon_template_contents = "No rimon Template Given by User."
        output_contents = []
        for file in self.files:
            file_extension = Path(file.name).suffix

            output_contents.append(f"\n\n{file.name} has the following contents:\n\n")

            match file_extension:
                case ".pdf":
                    contents = self.load_pdf(file)
                case ".txt":
                    contents = self.load_txt(file)
                case ".docx":
                    contents, TEMPLATE_FLAG = self.load_docx_advanced(file)
                case ".rtf":
                    contents = self.load_rtf(file)
                case ".csv":
                    contents = self.load_csv(file)
                case ".xlsx":
                    contents = self.load_xlsx(file)
                case _:
                    contents = "Unknown file type"

            if TEMPLATE_FLAG:
                rimon_template_contents = contents
                TEMPLATE_FLAG = False
            else:
                output_contents.append(contents)
            logger.info(f"Loaded {file.name} with file contents: {contents[:200]}")

        rimon_tokens = self.count_tokens(rimon_template_contents)
        additional_contents_tokens = self.count_tokens("".join(output_contents))

        return (
            output_contents,
            rimon_template_contents,
            rimon_tokens + additional_contents_tokens,
        )

    # TODO: To set up with pdf2text, I think we need to make a tempfile from the bytes object then read in that file
    def load_pdf(self, file_object=None, file_path=None) -> str:
        """
        Load a .pdf file, and return the contents as a string
        """
        if file_path:
            pdf_document = fitz.open(file_path, filetype="pdf")
            filename = file_path
        else:
            pdf_document = fitz.open(stream=file_object.getvalue(), filetype="pdf")
            filename = file_object.name

        n_pages = len(pdf_document)

        text_chunks = [
            f"{pdf_document.load_page(page_number).get_text()}"
            for page_number in range(n_pages)
        ]

        output_text = "".join(text_chunks)

        return output_text

    # def load_pdf_advanced(self, file_object) -> str:
    #     """
    #     Load a .pdf file, and return the contents as a string
    #     """

    #     with NamedTemporaryFile(suffix=".pdf", delete=True) as temp_pdf_file:
    #         temp_pdf_file.write(BytesIO(file_object.getvalue()).read())

    #         with open(temp_pdf_file, "rb") as f:
    #             pdf = pdftotext.PDF(f)

    #     return "\n\n".join(pdf)

    def load_rtf(self, file_object=None, file_path=None) -> str:
        """
        Load a .rtf file, and return the contents as a string
        """
        if file_path:
            with open(file_path, "r") as file:
                rtf_contents = file.read()
            output_text = rtf_to_text(rtf_contents)
            filename = file_path
        else:
            output_text = rtf_to_text(file_object.read().decode("utf-8"))
            filename = file_object.name

        return output_text

    def load_txt(self, file_object=None, file_path=None) -> str:
        """
        Load a .txt file, and return the contents as a string
        """
        if file_path:
            with open(file_path, "r") as file:
                output_text = file.read()
            filename = file_path
        else:
            output_text = StringIO(file_object.getvalue().decode("utf-8")).read()
            filename = file_object.name

        return output_text

    def load_docx(self, file_object=None, file_path=None) -> str:
        """
        Load a .docx file, and return the contents as a string
        """
        if file_path:
            doc = docx.Document(file_path)
            filename = file_path
        else:
            docx_file = BytesIO(file_object.getvalue())
            doc = docx.Document(docx_file)
            filename = file_object.name

        output_text = [f"{para.text}\n" for para in doc.paragraphs]

        return "".join(output_text)

    def load_docx_advanced(self, file_object=None, file_path=None) -> str:
        """
        Load a .docx file, and return the contents as a string
        """

        if file_path:
            output_text = docx2txt.process(file_path)
            return output_text
        else:
            TEMPLATE_FLAG = False
            with NamedTemporaryFile(suffix=".docx", delete=True) as temp_docx_file:
                temp_docx_file.write(BytesIO(file_object.getvalue()).read())
                doc_text = docx2txt.process(temp_docx_file.name)

                # Now, we need to check the doc_text to see if it was actually an rimon template!
                # If it was, we will instead call
                if (
                    "Please fill out the following table:" in doc_text
                    or "What was your goal or problem being solved in this Activity?"
                    in doc_text
                ):
                    doc_text = parse_rimon_template(file_object)
                    TEMPLATE_FLAG = True

            return doc_text, TEMPLATE_FLAG

    def load_xlsx(self, file_object=None, file_path=None) -> str:
        """
        Load a .xlsx file, and return the contents as a string
        """
        if file_path:
            xlsx_file = pd.ExcelFile(file_path)
            filename = file_path
        else:
            xlsx_file = pd.ExcelFile(StringIO(file_object.getvalue().decode("latin-1")))
            filename = file_object.name

        xlsx_file_sheets = {
            sheet_name: xlsx_file.parse(sheet_name)
            for sheet_name in xlsx_file.sheet_names
        }

        output_contents = []
        for xlsx_sheet_name, xlsx_sheet_contents in xlsx_file_sheets.items():
            output_contents.append(
                f"\n\nThe Sheet: {xlsx_sheet_name} has the following contents:\n"
            )
            output_contents.append(xlsx_sheet_contents.to_string(index=False))

        output_text = "".join(output_contents)

        return output_text

    def load_csv(self, file_object=None, file_path=None) -> str:
        """
        Load a .csv file, and return the contents as a string
        """
        if file_path:
            data = pd.read_csv(file_path)
            filename = file_path
        else:
            data = pd.read_csv(StringIO(file_object.getvalue().decode("utf-8")))
            filename = file_object.name

        output_text = data.to_string(index=False)
        return output_text


def parse_rimon_template(
    INPUT_DOCX, TEMPLATE_PATH=NEW_TEMPLATE_PATH, OLD_TEMPLATE_PATH=OLD_TEMPLATE_PATH
) -> str:
    LINES_TO_KEEP = {
        "Company name",
        "Company postcode",
        "Project name",
        "Project start date",
        "Estimated end date",
        "1b) Project Objectives:",
        "1c) Project Activities:",
        "Step 2 – Core Activities",
        "Step 3 - Supporting Activities",
    }

    # This section can all probably be cached, so it doesn't need computed each time?
    template = docx2txt.process(TEMPLATE_PATH)

    template_lines = template.split("\n")

    while "" in template_lines:
        template_lines.remove("")

    template_lines.pop(1)

    for l in LINES_TO_KEEP:
        try:
            template_lines.remove(
                l
            )  # keep section titles to indicate section begginings
        except:
            continue

    template2 = docx2txt.process(OLD_TEMPLATE_PATH)

    template_lines2 = template2.split("\n")
    while "" in template_lines2:
        template_lines2.remove("")

    template_lines2.pop(1)  # line with financial year
    for l in LINES_TO_KEEP:
        try:
            template_lines2.remove(
                l
            )  # keep section titles to indicate section begginings
        except:
            continue

    template_lines = template_lines + template_lines2

    with NamedTemporaryFile(suffix=".docx", delete=True) as temp_docx_file:
        temp_docx_file.write(BytesIO(INPUT_DOCX.getvalue()).read())
        text = docx2txt.process(temp_docx_file.name)

    lines = text.split("\n")
    while "" in lines:
        lines.remove("")

    # Remove unnecessary text
    for line in template_lines:
        try:
            lines.remove(line)
        except:
            continue

    lines = [f"{line}\n" for line in lines]

    return "\n".join(lines)


def process_files(files, summarize: bool = False, **runpod_credentials):
    """
    Takes streamlit files, and loads the document parser and simply extracts out the file contents.
    """

    parser = DocumentParser(files)
    file_contents, rimon_summary, total_tokens = parser.load_file_contents()

    # Now, call summarization to summarize the file contents ...
    if summarize:
        file_contents_summarized = summarize_list_of_text(
            file_contents, **runpod_credentials
        )
        return (
            file_contents_summarized,
            rimon_summary,
            total_tokens,
        )
    else:

        return (
            "".join(file_contents),
            rimon_summary,
            total_tokens,
        )


def summarize_list_of_text(input_text_chunks: list[str], **runpod_credentials) -> str:
    """
    Take a list of strings, and summarize them into one string
    """
    encoding = tiktoken.get_encoding("cl100k_base")
    output = []

    for index, input_text_chunk in enumerate(input_text_chunks):
        # Firstly, let's only pass in up to 4k tokens at a time. We will use a 500 token overlap
        input_text_chunk_tokens = encoding.encode(input_text_chunk)

        input_text_chunk_sub_chunks = [
            encoding.decode(input_text_chunk_tokens[i : i + 6000])
            for i in range(0, len(input_text_chunk_tokens), 5500)
        ]

        for sub_index, sub_chunk in enumerate(input_text_chunk_sub_chunks):
            input_prompt = summarization_prompt_rimon_specific_json_style.replace(
                "<text-to-compress>", sub_chunk
            )
            summary_output = runpod_call(input_prompt, **runpod_credentials)

            output.append(summary_output)

    return "\n\n".join(output)


def summarize_text(input_text: str, **runpod_credentials):
    """
    Summarize one text segment
    """
    input_prompt = summarization_prompt_rimon_specific.replace(
        "<text-to-compress>", input_text
    )

    return runpod_call(input_prompt, **runpod_credentials)
